"""Document styling and element rendering for docx review packets.

This module is the single place to iterate on visual appearance. It owns
both the style parameters (fonts, sizes, colors, spacing) and the methods
that convert markdown elements into python-docx objects. To change how the
document looks, edit this file — the structural orchestrator in ``docx.py``
should rarely need changes.

The :class:`DocumentStyle` class can be subclassed for radically different
looks, or the defaults can be tweaked for incremental refinement.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import Any

import mistune
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from clarity_agent.packet import PacketPart


class DocumentStyle:
    """Visual styling and element rendering for docx packets.

    This class owns two responsibilities:

    1. **Style parameters** — fonts, sizes, colors, spacing, margins.
       Adjust these attributes on an instance (or in ``__init__``) to
       change the look without touching rendering logic.

    2. **Element rendering methods** — convert structural and markdown
       elements into python-docx objects. Override individual methods
       to change how specific elements are rendered (e.g. how lists are
       indented or how code blocks are shaded).
    """

    def __init__(self) -> None:
        # Fonts.
        self.heading_font: str = "Calibri"
        self.body_font: str = "Calibri"
        self.mono_font: str = "Consolas"

        # Font sizes (Pt).
        self.title_size: float = 24
        self.part_title_size: float = 18
        self.section_heading_size: float = 15
        self.subsection_heading_size: float = 13
        self.body_size: float = 11

        # Colors (RGB tuples).
        self.heading_color: tuple[int, int, int] = (0, 51, 102)
        self.body_color: tuple[int, int, int] = (51, 51, 51)
        self.accent_color: tuple[int, int, int] = (0, 102, 153)
        self.code_bg_color: tuple[int, int, int] = (245, 245, 245)

        # Spacing (Pt).
        self.body_spacing_after: float = 8
        self.heading_spacing_before: float = 18
        self.heading_spacing_after: float = 6
        self.line_spacing: float = 1.15

        # Page.
        self.page_margin_inches: float = 1.0

        # Markdown parser (created once, reused).
        # Enable table plugin for markdown table support.
        self._md: Any = mistune.create_markdown(renderer="ast", plugins=["table"])

    # -----------------------------------------------------------------
    # Helpers
    # -----------------------------------------------------------------

    def _rgb(self, color: tuple[int, int, int]) -> RGBColor:
        return RGBColor(*color)

    def _setup_heading_style(
        self,
        doc: Document,
        style_name: str,
        *,
        font: str,
        size: float,
        color: tuple[int, int, int],
        bold: bool = True,
        space_before: float = 0,
        space_after: float = 0,
    ) -> None:
        """Configure a built-in heading style with our visual parameters."""
        style = doc.styles[style_name]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.color.rgb = self._rgb(color)
        style.font.bold = bold
        pf = style.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)

    def _style_run(
        self,
        run: Any,
        *,
        font: str | None = None,
        size: float | None = None,
        color: tuple[int, int, int] | None = None,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """Apply font properties to a run."""
        if font:
            run.font.name = font
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = self._rgb(color)
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    # -----------------------------------------------------------------
    # Document-level methods (called by the orchestrator)
    # -----------------------------------------------------------------

    def setup_page(self, doc: Document) -> None:
        """Configure page margins, default font, and heading styles."""
        for section in doc.sections:
            margin = Inches(self.page_margin_inches)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Set default font on the document style.
        style = doc.styles["Normal"]
        style.font.name = self.body_font
        style.font.size = Pt(self.body_size)
        style.font.color.rgb = self._rgb(self.body_color)
        pf = style.paragraph_format
        pf.space_after = Pt(self.body_spacing_after)
        pf.line_spacing = self.line_spacing

        # Configure built-in heading styles so that Word's navigation
        # pane, outline view, and automatic TOC work correctly.
        self._setup_heading_style(
            doc, "Title",
            font=self.heading_font, size=self.title_size,
            color=self.heading_color, space_after=4,
        )
        self._setup_heading_style(
            doc, "Heading 1",
            font=self.heading_font, size=self.part_title_size,
            color=self.heading_color,
            space_before=self.heading_spacing_before,
            space_after=self.heading_spacing_after,
        )
        self._setup_heading_style(
            doc, "Heading 2",
            font=self.heading_font, size=self.section_heading_size,
            color=self.heading_color,
            space_before=self.heading_spacing_before,
            space_after=self.heading_spacing_after,
        )
        self._setup_heading_style(
            doc, "Heading 3",
            font=self.heading_font, size=self.subsection_heading_size,
            color=self.heading_color,
            space_before=14, space_after=self.heading_spacing_after,
        )
        self._setup_heading_style(
            doc, "Heading 4",
            font=self.heading_font, size=self.subsection_heading_size - 1,
            color=self.heading_color,
            space_before=14, space_after=self.heading_spacing_after,
        )
        self._setup_heading_style(
            doc, "Heading 5",
            font=self.heading_font, size=self.body_size + 1,
            color=self.heading_color,
            space_before=12, space_after=self.heading_spacing_after,
        )

    def add_title(self, doc: Document, text: str) -> None:
        """Add the document title (using the built-in Title style)."""
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_toc(self, doc: Document, parts: list[PacketPart]) -> None:
        """Add a static table of contents grouped by part."""
        p = doc.add_paragraph()
        run = p.add_run("Contents")
        self._style_run(
            run,
            font=self.heading_font,
            size=self.section_heading_size,
            color=self.heading_color,
            bold=True,
        )
        p.paragraph_format.space_after = Pt(6)

        for part in parts:
            # Part title (bold line).
            p = doc.add_paragraph()
            run = p.add_run(part.title)
            self._style_run(
                run, font=self.body_font, size=self.body_size, bold=True,
            )
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)

            # Section titles (indented).
            for section in part.sections:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(section.title)
                self._style_run(
                    run,
                    font=self.body_font,
                    size=self.body_size,
                    color=self.accent_color,
                )

    def add_part_title(self, doc: Document, text: str) -> None:
        """Add a part title (e.g. 'Intro', 'Solution') using Heading 1."""
        p = doc.add_heading(text.upper(), level=1)
        # Add a thin line under the part title.
        _add_bottom_border(p)

    def add_section_heading(self, doc: Document, text: str) -> None:
        """Add a section heading (e.g. 'Problem Statement') using Heading 2."""
        doc.add_heading(text, level=2)

    def add_subsection_heading(self, doc: Document, text: str) -> None:
        """Add a subsection heading (e.g. 'Failure: SQL Injection') using Heading 3."""
        doc.add_heading(text, level=3)

    def add_separator(self, doc: Document) -> None:
        """Add a visual separator between parts (page break)."""
        doc.add_page_break()

    def add_footer(self, doc: Document) -> None:
        """Add a generation timestamp footer."""
        now = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
        # Thin rule.
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        _add_top_border(p)
        # Timestamp.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Generated {now}")
        self._style_run(
            run, font=self.body_font, size=9, color=(153, 153, 153), italic=True,
        )

    # -----------------------------------------------------------------
    # Markdown rendering (the core conversion engine)
    # -----------------------------------------------------------------

    def render_markdown(
        self, doc: Document, text: str, base_heading_level: int = 2,
    ) -> None:
        """Parse markdown text and append docx elements to *doc*.

        *base_heading_level* controls where headings in the markdown content
        map to in the document hierarchy. For example, a ``## Foo`` in the
        source with ``base_heading_level=2`` becomes a level-3 heading.
        """
        tokens: list[dict[str, Any]] = self._md(text)
        self._render_tokens(doc, tokens, base_heading_level)

    def _render_tokens(
        self,
        doc: Document,
        tokens: list[dict[str, Any]],
        base_heading_level: int,
    ) -> None:
        """Walk AST tokens and dispatch to element methods."""
        for token in tokens:
            tok_type: str = token["type"]

            if tok_type == "paragraph":
                self.render_paragraph(doc, token.get("children", []))
            elif tok_type == "heading":
                level: int = token["attrs"]["level"]
                # Shift heading levels relative to document structure.
                effective_level: int = base_heading_level + level - 1
                self.render_heading(
                    doc, effective_level, token.get("children", []),
                )
            elif tok_type == "list":
                ordered: bool = token["attrs"].get("ordered", False)
                items: list[dict[str, Any]] = token.get("children", [])
                self.render_list(doc, items, ordered)
            elif tok_type == "table":
                self.render_table(doc, token)
            elif tok_type == "block_code":
                code: str = token.get("raw", "")
                lang: str = token.get("attrs", {}).get("info", "")
                self.render_code_block(doc, code, lang)
            elif tok_type == "blank_line":
                pass  # Ignore blank lines in AST.

    # -----------------------------------------------------------------
    # Element-level methods (individually overridable)
    # -----------------------------------------------------------------

    def render_table(
        self, doc: Document, token: dict[str, Any],
    ) -> None:
        """Render a markdown table as a Word table."""
        children: list[dict[str, Any]] = token.get("children", [])
        if not children:
            return

        # Count columns from first row.
        first_row = children[0].get("children", [])
        num_cols = len(first_row)
        if num_cols == 0:
            return

        table = doc.add_table(rows=0, cols=num_cols)
        table.style = "Table Grid"

        for row_token in children:
            row_type = row_token.get("type", "")
            cells_data = row_token.get("children", [])

            if row_type == "table_head":
                row = table.add_row()
                for i, cell in enumerate(cells_data):
                    if i < num_cols:
                        text = _children_to_text(cell.get("children", []))
                        run = row.cells[i].paragraphs[0].add_run(text)
                        self._style_run(run, font=self.body_font, size=self.body_size, bold=True)
            elif row_type == "table_body":
                for body_row in cells_data:
                    body_cells = body_row.get("children", [])
                    row = table.add_row()
                    for i, cell in enumerate(body_cells):
                        if i < num_cols:
                            text = _children_to_text(cell.get("children", []))
                            run = row.cells[i].paragraphs[0].add_run(text)
                            self._style_run(run, font=self.body_font, size=self.body_size)

    def render_paragraph(
        self, doc: Document, children: list[dict[str, Any]],
    ) -> None:
        """Render a paragraph with inline children."""
        p = doc.add_paragraph()
        self.render_inline(p, children)

    def render_heading(
        self,
        doc: Document,
        level: int,
        children: list[dict[str, Any]],
    ) -> None:
        """Render a heading at the given document level using Word styles."""
        text: str = _children_to_text(children)
        # Clamp to valid Word heading levels (1–9).
        heading_level: int = min(max(level, 1), 9)
        doc.add_heading(text, level=heading_level)

    def render_list(
        self,
        doc: Document,
        items: list[dict[str, Any]],
        ordered: bool,
    ) -> None:
        """Render a list (ordered or unordered)."""
        for i, item in enumerate(items):
            children: list[dict[str, Any]] = item.get("children", [])
            # list_item children are typically block_text nodes.
            inline_children: list[dict[str, Any]] = []
            for child in children:
                if child["type"] == "block_text" or child["type"] == "paragraph":
                    inline_children.extend(child.get("children", []))

            bullet: str = f"{i + 1}." if ordered else "\u2022"
            self.render_list_item(doc, inline_children, bullet, indent_level=0)

    def render_list_item(
        self,
        doc: Document,
        children: list[dict[str, Any]],
        bullet: str,
        indent_level: int,
    ) -> None:
        """Render a single list item."""
        p = doc.add_paragraph()
        indent: float = 0.4 + (indent_level * 0.3)
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        # Bullet/number prefix.
        run = p.add_run(f"{bullet}  ")
        self._style_run(run, font=self.body_font, size=self.body_size)

        # Inline content.
        self.render_inline(p, children)

    def render_code_block(
        self, doc: Document, code: str, language: str,
    ) -> None:
        """Render a fenced code block with monospace font and background.

        Mermaid blocks attempt PNG rendering via mermaid-py (which calls
        the Mermaid.ink service). This is slow (~2-5s per diagram) so it
        can be disabled by setting ``self.render_mermaid_images = False``.
        """
        if language == "mermaid" and getattr(self, "render_mermaid_images", True):
            if self._render_mermaid_image(doc, code):
                return
            # Fall through to text rendering if image generation fails.

        for line in code.rstrip("\n").split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _set_paragraph_shading(p, self.code_bg_color)

            run = p.add_run(line)
            self._style_run(
                run, font=self.mono_font, size=self.body_size - 1,
                color=self.body_color,
            )

    def _render_mermaid_image(self, doc: Document, code: str) -> bool:
        """Try to render Mermaid code as a PNG image in the document.

        Uses ``mermaid-py`` to convert Mermaid syntax to PNG via the
        Mermaid.ink service. Falls back to text rendering if unavailable.
        """
        try:
            import os
            import tempfile

            import mermaid as md
            from mermaid.graph import Graph

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name

            g = Graph("threat_model", code)
            rendered = md.Mermaid(g)
            rendered.to_png(tmp_path)

            if os.path.getsize(tmp_path) > 0:
                doc.add_picture(tmp_path, width=Inches(5.5))
                os.unlink(tmp_path)
                return True

            os.unlink(tmp_path)
            return False
        except Exception:
            return False

    def render_inline(
        self, paragraph: Any, children: list[dict[str, Any]],
    ) -> None:
        """Walk inline tokens and append runs to a paragraph."""
        for child in children:
            child_type: str = child["type"]

            if child_type == "text":
                run = paragraph.add_run(child.get("raw", ""))
                self._style_run(
                    run, font=self.body_font, size=self.body_size,
                    color=self.body_color,
                )
            elif child_type == "strong":
                self._render_inline_styled(
                    paragraph, child.get("children", []), bold=True,
                )
            elif child_type == "emphasis":
                self._render_inline_styled(
                    paragraph, child.get("children", []), italic=True,
                )
            elif child_type == "codespan":
                run = paragraph.add_run(child.get("raw", ""))
                self._style_run(
                    run, font=self.mono_font, size=self.body_size - 1,
                    color=self.accent_color,
                )
            elif child_type == "link":
                link_text: str = _children_to_text(child.get("children", []))
                url: str = child.get("attrs", {}).get("url", "")
                self.render_link(paragraph, link_text, url)
            elif child_type == "softbreak":
                paragraph.add_run("\n")

    def _render_inline_styled(
        self,
        paragraph: Any,
        children: list[dict[str, Any]],
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """Render inline children with bold/italic applied."""
        for child in children:
            text: str = child.get("raw", "")
            if child["type"] == "text":
                run = paragraph.add_run(text)
                self._style_run(
                    run, font=self.body_font, size=self.body_size,
                    color=self.body_color, bold=bold, italic=italic,
                )
            elif child["type"] == "codespan":
                run = paragraph.add_run(text)
                self._style_run(
                    run, font=self.mono_font, size=self.body_size - 1,
                    color=self.accent_color, bold=bold, italic=italic,
                )

    def render_link(
        self, paragraph: Any, text: str, url: str,
    ) -> None:
        """Render a link as styled text with URL."""
        run = paragraph.add_run(text)
        self._style_run(
            run, font=self.body_font, size=self.body_size,
            color=self.accent_color,
        )
        run.underline = True


# =====================================================================
# Module-level helpers (XML manipulation for python-docx)
# =====================================================================

def _children_to_text(children: list[dict[str, Any]]) -> str:
    """Flatten inline AST children to plain text."""
    parts: list[str] = []
    for child in children:
        if "raw" in child:
            parts.append(child["raw"])
        elif "children" in child:
            parts.append(_children_to_text(child["children"]))
    return "".join(parts)


def _add_bottom_border(paragraph: Any) -> None:
    """Add a thin bottom border to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border(paragraph: Any) -> None:
    """Add a thin top border to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    top = pBdr.makeelement(
        qn("w:top"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        },
    )
    pBdr.append(top)
    pPr.append(pBdr)


def _set_paragraph_shading(
    paragraph: Any, color: tuple[int, int, int],
) -> None:
    """Set background shading on a paragraph via XML."""
    hex_color: str = "{:02X}{:02X}{:02X}".format(*color)
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): hex_color,
        },
    )
    pPr.append(shd)


DEFAULT_STYLE = DocumentStyle()
