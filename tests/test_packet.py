"""Tests for the packet generator."""

from __future__ import annotations

from pathlib import Path
from typing import cast

import pytest

from clarity_agent.packet import (
    ContentSource,
    DirectorySource,
    Packet,
    PacketContent,
    PacketError,
    SingleFileSource,
    generate_packet,
    register_format,
    register_source,
)
from clarity_agent.protocol.initialize import init_protocol

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_protocol(tmp_path: Path, docs: dict[str, str] | None = None) -> Path:
    """Create a .clarity-protocol/ with optional real content."""
    init_protocol(tmp_path)
    pd = tmp_path / ".clarity-protocol"
    if docs:
        for rel_path, content in docs.items():
            (pd / rel_path).parent.mkdir(parents=True, exist_ok=True)
            (pd / rel_path).write_text(content)
    return pd


# ---------------------------------------------------------------------------
# SingleFileSource
# ---------------------------------------------------------------------------

class TestSingleFileSource:
    """SingleFileSource reads a single protocol document."""

    def test_collects_real_content(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nWe need faster coffee.\n",
        })
        source = SingleFileSource("Problem Statement", "goal/problem.md")
        result = source.collect(pd)
        assert result is not None
        assert result.title == "Problem Statement"
        assert result.intro is not None
        assert "faster coffee" in result.intro
        assert result.sections is None

    def test_skips_template(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path)
        # problem.md is a template after init_protocol
        source = SingleFileSource("Problem Statement", "goal/problem.md")
        assert source.collect(pd) is None

    def test_skips_missing_file(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path)
        source = SingleFileSource("Nonexistent", "goal/nonexistent.md")
        assert source.collect(pd) is None


# ---------------------------------------------------------------------------
# DirectorySource
# ---------------------------------------------------------------------------

class TestDirectorySource:
    """DirectorySource reads an index + individual files from a directory."""

    def test_collects_index_and_files(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failures.md": "# Failure Modes\n\n2 failures identified.\n",
            "failures/failure-01-sql-injection.md": (
                "# Failure: SQL Injection\n\n## Summary\n\nBad input.\n"
            ),
            "failures/failure-02-auth-bypass.md": (
                "# Failure: Auth Bypass\n\n## Summary\n\nWeak auth.\n"
            ),
        })
        source = DirectorySource("Failure Modes", "failures", "failures.md")
        result = source.collect(pd)
        assert result is not None
        assert result.title == "Failure Modes"
        assert result.intro is not None
        assert "2 failures" in result.intro
        assert result.sections is not None
        assert len(result.sections) == 2

    def test_skips_template_index_but_collects_files(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            # failures.md is still a template, but we add a real failure file
            "failures/failure-01-real.md": "# Failure: Real Problem\n\nBad stuff.\n",
        })
        source = DirectorySource("Failure Modes", "failures", "failures.md")
        result = source.collect(pd)
        assert result is not None
        assert result.intro is None  # template index skipped
        assert result.sections is not None
        assert len(result.sections) == 1

    def test_extracts_title_from_heading(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failure-01-injection.md": "# Failure: SQL Injection\n\nDetails.\n",
        })
        source = DirectorySource("Failure Modes", "failures", "failures.md")
        result = source.collect(pd)
        assert result is not None
        assert result.sections is not None
        assert "Failure: SQL Injection" in result.sections

    def test_falls_back_to_stem_without_heading(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failure-01-no-heading.md": "Just some text without a heading.\n",
        })
        source = DirectorySource("Failure Modes", "failures", "failures.md")
        result = source.collect(pd)
        assert result is not None
        assert result.sections is not None
        assert "failure-01-no-heading" in result.sections

    def test_empty_dir_returns_none(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path)
        # failures dir has only the template index
        source = DirectorySource("Failure Modes", "failures", "failures.md")
        assert source.collect(pd) is None

    def test_missing_dir_returns_none(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path)
        source = DirectorySource("Stuff", "nonexistent", "index.md")
        assert source.collect(pd) is None

    def test_include_index_false_skips_index(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failures.md": "# Failure Modes\n\n2 failures identified.\n",
            "failures/failure-01-sql-injection.md": (
                "# Failure: SQL Injection\n\n## Summary\n\nBad input.\n"
            ),
        })
        source = DirectorySource("Failure Analysis", "failures", "failures.md", include_index=False)
        result = source.collect(pd)
        assert result is not None
        assert result.intro is None  # Index was excluded
        assert result.sections is not None
        assert len(result.sections) == 1

    def test_include_index_true_is_default(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failures.md": "# Failure Modes\n\n2 failures identified.\n",
            "failures/failure-01-sql-injection.md": (
                "# Failure: SQL Injection\n\nBad input.\n"
            ),
        })
        source = DirectorySource("Failure Modes", "failures", "failures.md")
        result = source.collect(pd)
        assert result is not None
        assert result.intro is not None  # Index included by default


# ---------------------------------------------------------------------------
# Registration
# ---------------------------------------------------------------------------

class TestRegistration:
    """Source and format registration."""

    def test_register_custom_source(self, tmp_path: Path) -> None:
        """A custom ContentSource can be registered and used."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nReal content.\n",
        })

        class ConstantSource(ContentSource):
            def collect(self, protocol_dir: Path) -> PacketContent | None:
                return PacketContent(title=self.title, intro="Custom content here.")

        register_source("custom-test", ConstantSource("Custom"))
        try:
            output = generate_packet(pd, include=["custom-test"]).decode()
            assert "Custom content here." in output
        finally:
            # Clean up to avoid polluting other tests.
            from clarity_agent.packet import _SOURCES
            _SOURCES.pop("custom-test", None)

    def test_list_sources_with_titles(self) -> None:
        """list_sources_with_titles returns {id, title} pairs with correct titles."""
        from clarity_agent.packet import list_sources_with_titles

        result = list_sources_with_titles()
        assert isinstance(result, list)
        assert all("id" in entry and "title" in entry for entry in result)
        oq = next(e for e in result if e["id"] == "openquestions")
        assert oq["title"] == "Open Questions"

    def test_list_parts(self) -> None:
        """list_parts returns canonical part structure."""
        from clarity_agent.packet import list_parts

        result = list_parts()
        assert isinstance(result, list)
        assert all("title" in p and "source_ids" in p for p in result)
        first_part = result[0]
        # list_parts is typed as list[dict[str, object]] so we have to
        # narrow the heterogeneous-value dict back to the concrete
        # list type before doing membership checks.
        first_source_ids = cast(list[str], first_part["source_ids"])
        assert "problem" in first_source_ids

    def test_unknown_source_raises(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path)
        with pytest.raises(PacketError, match="Unknown content source"):
            generate_packet(pd, include=["nonexistent"])

    def test_unknown_format_raises(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nReal.\n",
        })
        with pytest.raises(PacketError, match="Unknown format"):
            generate_packet(pd, include=["problem"], format="pdf")

    def test_register_custom_format(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nSome problem.\n",
        })

        def plain_text_renderer(packet: Packet) -> bytes:
            return "\n".join(
                section.title
                for part in packet.parts
                for section in part.sections
            ).encode("utf-8")

        register_format("plain-test", plain_text_renderer)
        try:
            output = generate_packet(pd, include=["problem"], format="plain-test").decode()
            assert output == "Problem Statement"
        finally:
            from clarity_agent.packet import _FORMAT_REGISTRY
            _FORMAT_REGISTRY.pop("plain-test", None)


# ---------------------------------------------------------------------------
# generate_packet
# ---------------------------------------------------------------------------

class TestGeneratePacket:
    """generate_packet() assembles content from sources."""

    def test_single_file_categories(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nFaster coffee.\n",
            "solution/solution.md": "# Solution\n\nBigger kettle.\n",
        })
        output = generate_packet(pd, include=["problem", "solution"]).decode()
        assert "Faster coffee" in output
        assert "Bigger kettle" in output

    def test_directory_category(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failures.md": "# Failure Modes\n\n1 failure.\n",
            "failures/failure-01-injection.md": (
                "# Failure: SQL Injection\n\nBad input.\n"
            ),
        })
        # Index source gets the overview
        output_index = generate_packet(pd, include=["failures_index"]).decode()
        assert "1 failure" in output_index
        # Detail source gets the individual failures
        output_detail = generate_packet(pd, include=["failures_detail"]).decode()
        assert "SQL Injection" in output_detail

    def test_skips_empty_sources(self, tmp_path: Path) -> None:
        """All-template protocol raises PacketError."""
        pd = _make_protocol(tmp_path)
        with pytest.raises(PacketError, match="No content collected"):
            generate_packet(pd)

    def test_include_none_collects_all_non_template(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nReal problem.\n",
            "goal/stakeholders.md": "# Stakeholders\n\n- Users\n",
            # requirements left as template
        })
        output = generate_packet(pd, include=None).decode()
        assert "Real problem" in output
        assert "Users" in output
        # Template requirements should not appear
        assert "[To be determined" not in output

    def test_title_from_summary(self, tmp_path: Path) -> None:
        """Document title comes from summary.md heading."""
        pd = _make_protocol(tmp_path, {
            "summary.md": "# CoffeeFast\n\nA faster coffee brewing system.\n",
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        output = generate_packet(pd, include=["problem"]).decode()
        assert "# CoffeeFast" in output
        assert "Review Packet" not in output

    def test_summary_as_preamble(self, tmp_path: Path) -> None:
        """Summary body appears before the table of contents."""
        pd = _make_protocol(tmp_path, {
            "summary.md": "# CoffeeFast\n\nA faster coffee brewing system.\n",
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        output = generate_packet(pd, include=["problem"]).decode()
        summary_pos = output.index("faster coffee brewing")
        toc_pos = output.index("## Contents")
        assert summary_pos < toc_pos

    def test_fallback_title_without_summary(self, tmp_path: Path) -> None:
        """Without summary.md, title falls back to 'Review Packet'."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            # summary.md left as template
        })
        output = generate_packet(pd, include=["problem"]).decode()
        assert "# Review Packet" in output

    def test_summary_always_included(self, tmp_path: Path) -> None:
        """Summary appears even when not in include list."""
        pd = _make_protocol(tmp_path, {
            "summary.md": "# CoffeeFast\n\nBrewing system.\n",
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        # include only mentions "problem", not "summary"
        output = generate_packet(pd, include=["problem"]).decode()
        assert "# CoffeeFast" in output
        assert "Brewing system" in output


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------

class TestMarkdownRenderer:
    """The markdown renderer produces legible documents."""

    def test_has_toc(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "goal/stakeholders.md": "# Stakeholders\n\n- Alice\n",
        })
        output = generate_packet(pd, include=["problem", "stakeholders"]).decode()
        assert "## Contents" in output
        assert "Problem Statement" in output
        assert "Stakeholders" in output

    def test_sections_use_h2(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        output = generate_packet(pd, include=["problem"]).decode()
        assert "## Problem Statement" in output

    def test_subsections_use_h3(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failure-01-injection.md": (
                "# Failure: SQL Injection\n\nDetails.\n"
            ),
        })
        output = generate_packet(pd, include=["failures_detail"]).decode()
        assert "## Failure Analysis" in output
        assert "### Failure: SQL Injection" in output

    def test_index_before_detail_in_output(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failures.md": "# Failure Modes\n\nIndex overview.\n",
            "failures/failure-01-x.md": "# Failure: X\n\nDetails.\n",
        })
        output = generate_packet(pd, include=["failures_index", "failures_detail"]).decode()
        idx_pos = output.index("Index overview")
        sec_pos = output.index("### Failure: X")
        assert idx_pos < sec_pos

    def test_has_footer(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        output = generate_packet(pd, include=["problem"]).decode()
        assert "*Generated " in output
        assert "UTC*" in output

    def test_toc_grouped_by_part(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "goal/open-questions.md": "# Open Questions\n\nAre there any?\n",
        })
        output = generate_packet(pd, include=["problem", "openquestions"]).decode()
        assert "**The Problem and the Solution**" in output

    def test_parts_separated_by_rules(self, tmp_path: Path) -> None:
        """Different parts have horizontal rules between them."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "failures/failure-01-x.md": "# Failure: X\n\nOverview.\n",
        })
        output = generate_packet(pd, include=["problem", "failures_detail"]).decode()
        prob_pos = output.index("## Problem Statement")
        fail_pos = output.index("## Failure Analysis")
        between = output[prob_pos:fail_pos]
        assert "---" in between


# ---------------------------------------------------------------------------
# Canonical ordering
# ---------------------------------------------------------------------------

class TestCanonicalOrdering:
    """Sources are grouped into parts in a canonical order."""

    def test_intro_before_design_before_failures(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failure-01-x.md": "# Failure: X\n\nOverview.\n",
            "goal/problem.md": "# Problem\n\nContent.\n",
            "solution/solution.md": "# Solution\n\nContent.\n",
        })
        output = generate_packet(pd).decode()
        prob_pos = output.index("## Problem Statement")
        sol_pos = output.index("## Solution")
        fail_pos = output.index("## Failure Analysis")
        assert prob_pos < sol_pos < fail_pos

    def test_custom_source_goes_to_additional(self, tmp_path: Path) -> None:
        """Sources not in _CANONICAL_PARTS are placed after canonical ones."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })

        class ExtraSource(ContentSource):
            def collect(self, protocol_dir: Path) -> PacketContent | None:
                return PacketContent(title=self.title, intro="Extra content.")

        register_source("extra-test", ExtraSource("Extra"))
        try:
            output = generate_packet(pd, include=["problem", "extra-test"]).decode()
            prob_pos = output.index("## Problem Statement")
            extra_pos = output.index("## Extra")
            assert prob_pos < extra_pos
            assert "**Additional**" in output
        finally:
            from clarity_agent.packet import _SOURCES
            _SOURCES.pop("extra-test", None)

    def test_multiple_sources_in_same_part(self, tmp_path: Path) -> None:
        """Problem and open questions both appear under the same part."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "goal/open-questions.md": "# Open Questions\n\nAre there any?\n",
        })
        output = generate_packet(pd, include=["problem", "openquestions"]).decode()
        # Both under the same part, no separator between them
        toc = output[:output.index("---")]
        assert "**The Problem and the Solution**" in toc
        assert "Problem Statement" in toc
        assert "Open Questions" in toc


# ---------------------------------------------------------------------------
# Docx renderer
# ---------------------------------------------------------------------------

docx_mod = pytest.importorskip("docx", reason="python-docx not installed")

import io  # noqa: E402 — after importorskip guard

from docx import Document as DocxDocument  # noqa: E402


def _docx_paragraphs(data: bytes) -> list[str]:
    """Extract paragraph texts from docx bytes."""
    doc = DocxDocument(io.BytesIO(data))
    return [p.text for p in doc.paragraphs if p.text.strip()]


class TestDocxRenderer:
    """The docx renderer produces valid Word documents."""

    def test_produces_valid_docx(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        data = generate_packet(pd, include=["problem"], format="docx")
        # Should be a valid zip (docx is zip-based).
        assert data[:2] == b"PK"
        # Should open without error.
        doc = DocxDocument(io.BytesIO(data))
        assert len(doc.paragraphs) > 0

    def test_has_title(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        texts = _docx_paragraphs(
            generate_packet(pd, include=["problem"], format="docx"),
        )
        assert "Review Packet" in texts

    def test_section_headings_present(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "solution/solution.md": "# Solution\n\nDetails.\n",
        })
        texts = _docx_paragraphs(
            generate_packet(pd, include=["problem", "solution"], format="docx"),
        )
        assert "Problem Statement" in texts
        assert "Solution" in texts

    def test_subsection_headings_present(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "failures/failure-01-injection.md": (
                "# Failure: SQL Injection\n\nBad input.\n"
            ),
        })
        texts = _docx_paragraphs(
            generate_packet(pd, include=["failures_detail"], format="docx"),
        )
        assert "Failure: SQL Injection" in texts

    def test_content_rendered(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nWe need faster coffee.\n",
        })
        texts = _docx_paragraphs(
            generate_packet(pd, include=["problem"], format="docx"),
        )
        assert any("faster coffee" in t for t in texts)

    def test_bold_and_italic_rendered(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nThis is **bold** and *italic*.\n",
        })
        data = generate_packet(pd, include=["problem"], format="docx")
        doc = DocxDocument(io.BytesIO(data))
        # Find the paragraph with "bold" and check run formatting.
        for p in doc.paragraphs:
            runs_with_bold = [r for r in p.runs if r.bold and "bold" in r.text]
            if runs_with_bold:
                break
        else:
            pytest.fail("No bold run found containing 'bold'")

    def test_parts_separated_by_page_breaks(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "failures/failure-01-x.md": "# Failure: X\n\nOverview.\n",
        })
        data = generate_packet(pd, include=["problem", "failures_detail"], format="docx")
        doc = DocxDocument(io.BytesIO(data))
        # Look for a page break element between the two parts.
        has_page_break = any(
            p.paragraph_format.page_break_before
            for p in doc.paragraphs
        )
        # Also check for explicit page break runs.
        has_break_run = any(
            "lastRenderedPageBreak" in p._p.xml or "w:br" in p._p.xml
            for p in doc.paragraphs
        )
        assert has_page_break or has_break_run

    def test_has_footer(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        texts = _docx_paragraphs(
            generate_packet(pd, include=["problem"], format="docx"),
        )
        assert any("Generated" in t and "UTC" in t for t in texts)

    def test_code_block_uses_mono_font(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": (
                "# Problem\n\nSee code:\n\n```\nfoo = 42\n```\n"
            ),
        })
        data = generate_packet(pd, include=["problem"], format="docx")
        doc = DocxDocument(io.BytesIO(data))
        mono_runs = [
            r for p in doc.paragraphs for r in p.runs
            if r.font.name == "Consolas" and "foo" in r.text
        ]
        assert len(mono_runs) >= 1

    def test_list_items_rendered(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\n- Alpha\n- Beta\n",
        })
        texts = _docx_paragraphs(
            generate_packet(pd, include=["problem"], format="docx"),
        )
        assert any("Alpha" in t for t in texts)
        assert any("Beta" in t for t in texts)

    def test_heading_styles_used(self, tmp_path: Path) -> None:
        """Headings use actual Word heading styles for navigation/outline."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "failures/failure-01-injection.md": (
                "# Failure: SQL Injection\n\nBad input.\n"
            ),
        })
        data = generate_packet(
            pd, include=["problem", "failures_detail"], format="docx",
        )
        doc = DocxDocument(io.BytesIO(data))
        styles_used: dict[str, list[str]] = {}
        for p in doc.paragraphs:
            if p.text.strip():
                # python-docx types both p.style and style.name as
                # optional, but in practice the docx library always
                # populates them for paragraphs we care about.  Skip
                # safely if either is missing rather than asserting.
                if p.style is None or p.style.name is None:
                    continue
                styles_used.setdefault(p.style.name, []).append(p.text)

        # Title style for the document title.
        assert "Title" in styles_used
        # Heading 1 for part titles.
        assert "Heading 1" in styles_used
        # Heading 2 for section headings.
        assert "Heading 2" in styles_used
        assert any("Problem Statement" in t for t in styles_used["Heading 2"])
        # Heading 3 for subsection headings.
        assert "Heading 3" in styles_used
        assert any("SQL Injection" in t for t in styles_used["Heading 3"])


# ---------------------------------------------------------------------------
# Packet views
# ---------------------------------------------------------------------------

class TestPacketViews:
    """Packet views provide audience-specific document ordering."""

    def test_builtin_views_registered(self) -> None:
        from clarity_agent.packet import _VIEWS

        assert "complete" in _VIEWS
        assert "reviewer" in _VIEWS
        assert "product_manager" in _VIEWS
        assert "program_manager" in _VIEWS
        assert "engineer" in _VIEWS

    def test_complete_view_matches_canonical(self) -> None:
        from clarity_agent.packet import _CANONICAL_PARTS, _VIEWS

        assert _VIEWS["complete"].parts == list(_CANONICAL_PARTS)

    def test_view_ordering_applied(self, tmp_path: Path) -> None:
        """Engineer view puts problem context before architecture."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "solution/architecture.md": "# Architecture\n\nDesign.\n",
        })
        output = generate_packet(pd, view="engineer").decode()
        arch_pos = output.index("## Architecture")
        prob_pos = output.index("## Problem Statement")
        assert prob_pos < arch_pos

    def test_view_filters_sources(self, tmp_path: Path) -> None:
        """Engineer view does not include messaging."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "messaging.md": "# Messaging\n\nAudience info.\n",
        })
        output = generate_packet(pd, view="engineer").decode()
        assert "Audience info" not in output

    def test_view_with_include_intersection(self, tmp_path: Path) -> None:
        """include further filters within the view's sources."""
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
            "solution/architecture.md": "# Architecture\n\nDesign.\n",
        })
        output = generate_packet(pd, view="engineer", include=["problem"]).decode()
        assert "Content." in output
        assert "Design." not in output

    def test_unknown_view_raises(self, tmp_path: Path) -> None:
        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nContent.\n",
        })
        with pytest.raises(PacketError, match="Unknown view"):
            generate_packet(pd, view="nonexistent")

    def test_list_views_with_parts(self) -> None:
        from clarity_agent.packet import list_views_with_parts

        result = list_views_with_parts()
        assert isinstance(result, list)
        assert len(result) >= 5
        complete = next(v for v in result if v["id"] == "complete")
        assert "parts" in complete
        assert isinstance(complete["parts"], list)
        assert all(
            "title" in p and "source_ids" in p
            for p in complete["parts"]
        )


# ---------------------------------------------------------------------------
# Source size estimation
# ---------------------------------------------------------------------------

class TestSourceSizeEstimation:
    """estimate_source_sizes returns word counts for populated sources."""

    def test_returns_word_counts(self, tmp_path: Path) -> None:
        from clarity_agent.packet import estimate_source_sizes

        pd = _make_protocol(tmp_path, {
            "goal/problem.md": "# Problem\n\nWe need a faster way to brew coffee at scale.\n",
        })
        sizes = estimate_source_sizes(pd)
        assert "problem" in sizes
        assert sizes["problem"] > 0

    def test_omits_template_only_sources(self, tmp_path: Path) -> None:
        from clarity_agent.packet import estimate_source_sizes

        # Init protocol creates template content; don't write real docs.
        pd = _make_protocol(tmp_path)
        sizes = estimate_source_sizes(pd)
        # All files are templates, so nothing should have content.
        assert len(sizes) == 0

    def test_nonexistent_dir(self, tmp_path: Path) -> None:
        from clarity_agent.packet import estimate_source_sizes

        sizes = estimate_source_sizes(tmp_path / "nonexistent")
        assert sizes == {}
